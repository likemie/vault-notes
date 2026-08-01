from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path("raw/香港国民教育调研.docx")
OUTPUT = Path("raw/香港国民教育调研（修订版）.docx")


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def insert_before(paragraph, text="", style=None):
    result = paragraph.insert_paragraph_before(text)
    if style:
        result.style = style
    return result


def set_run_font(run, latin="Arial", east_asia="Songti SC", size=None, bold=None, color=None):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def set_all_xml_run_fonts(document):
    for run_el in document.element.body.iter(qn("w:r")):
        rpr = run_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_el.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:eastAsia"), "Songti SC")
        rfonts.set(qn("w:cs"), "Arial")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("香港国民教育调研（修订版）  ·  ")
    set_run_font(run, size=8, east_asia="PingFang SC", color=(100, 100, 100))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "PingFang SC")
    rpr.append(rfonts)
    field_run.append(rpr)
    field_run.extend([fld_begin, instr, fld_separate, page_text, fld_end])
    paragraph._p.append(field_run)


doc = Document(SOURCE)
paras = list(doc.paragraphs)

# Title and opening: clarify scope and avoid treating policy intent as observed effect.
set_text(
    paras[0],
    "香港中小学国民教育发展现状、实施张力与优化建议"
    "——基于政策文本、教材研究与前线教师材料的综合分析（截至2026年7月）",
)
set_text(
    paras[1],
    "近年香港基础教育中的国民教育进入制度化推进阶段。2024年4月“爱国主义教育工作小组”成立，"
    "2025年《香港国家安全教育课程框架》更新，小学人文科、初中公民、经济与社会科以及高中公民与"
    "社会发展科逐步形成跨学段课程链条。本报告综合政策文件、公开统计、媒体资料、三篇期刊论文、"
    "raw目录中的三篇相关中文论文以及三篇2026年硕士学位论文，分析课程、教材、师资和内地考察的"
    "进展与实施张力。报告所称“成效”主要指制度覆盖与课程建设进展；除非有直接数据，不将政策意图、"
    "教材内容或教师观察等同于学生学习结果和身份认同变化。",
)

scope_h = insert_after(paras[1], "研究范围与证据说明", "Heading 2")
scope_p = insert_after(
    scope_h,
    "本报告采用“政策设计—课程与教材—学校实施—学生经验”四层分析框架。官方文件用于确认制度安排；"
    "Vickers（2024）、Xu（2024）和Yan与Morris（2025）用于解释课程变革及教材叙事；彭俊源（2026）、"
    "杜雨璐（2026）和严瑶（2026）分别提供课程衔接、内地考察和教材内容的探索性证据。三篇学位论文"
    "具有材料新、问题聚焦的优势，但两篇属于文本分析，一篇采用目的性抽样访谈，均不能单独代表全港"
    "学校或证明因果关系。因此，正文在引用其发现时同时标明研究对象与适用边界。",
    "Normal",
)

set_text(paras[2], "一、香港中小学国民教育的制度进展与实施张力")
set_text(
    paras[3],
    "香港基础教育已建立贯穿小学、初中和高中的国民教育课程架构，课程指引、教材审核、教师培训与"
    "内地考察亦逐步配套。制度覆盖的扩大是明确进展，但课程的学段衔接、评估激励、教师专业能力和"
    "体验学习质量仍需分别检验，不能仅凭活动数量或文件完成度判断教育效果。",
)
set_text(paras[4], "（一）教学内容与课程设计：体系化建构与评估激励不足")
set_text(
    paras[5],
    "在宏观架构上，教育局已形成小学人文科、初中公民、经济与社会科、高中公民与社会发展科相互衔接"
    "的课程矩阵，并把宪法、基本法、一国两制、国家安全、中华文化和社会发展等内容纳入不同学段。",
)
set_text(
    paras[7],
    "高中公民科采用“达标/未达标”评级，2024年首届文凭试达标率为94%（见原稿资料10、13、15）。"
    "媒体访问中出现学生“无需温习”、学校挪用课时等说法，提示该科可能面临评估激励不足；但现有材料"
    "不足以推断这是全港普遍现象。更具结构性的证据来自彭俊源（2026）对初高中课程指引的编码比较："
    "高中高阶认知目标占32.2%，低于初中的41.4%，且个人发展内容在高中明显减少。由此，问题不仅是"
    "学生是否重视，更涉及课程目标和评价方式能否持续支持分析、评价、创造等高阶学习。",
)

new_h = insert_after(paras[7], "（二）课程衔接与认知进阶：选择性连续与目标错位", "Heading 3")
new_p1 = insert_after(
    new_h,
    "彭俊源（2026）以两份官方课程指引为对象，对初中116条、高中90条认知过程编码进行比较，并辅以"
    "文本分析。研究认为，两学段在宪制政治与国民身份议题上呈现“选择性连续”，但个人发展、程序性知识"
    "及高阶认知要求没有形成稳定递进；国民身份目标则更像框架替换，而非由近及远的渐进深化。这一发现"
    "表明，跨学段建设不能只核对主题是否重复，还应检验概念复杂度、探究工具和情感目标是否逐级发展。",
    "Normal",
)
new_p2 = insert_after(
    new_p1,
    "上述结论针对“意图课程”，不能直接说明教师实际教学或学生能力发生倒退；其价值在于为课程修订提供"
    "可检验的诊断假设。后续应以课堂任务、校本评估样本和学生作品验证课程文件中的认知要求是否真正转化"
    "为学习机会。",
    "Normal",
)

set_text(paras[8], "（三）教材编写与历史教育：叙事重构、内容覆盖与学术质量")
set_text(
    paras[9],
    "历史与教材是国民教育的重要载体。教育局于2025年公布优化高中中国历史科及历史科课程框架，拟于"
    "2027/28学年从中四级推行，并适用于2030年及以后的文凭试。相关变化应同时从课程覆盖、叙事选择、"
    "学术准确性和学生理解四个维度评价。",
)
set_text(
    paras[10],
    "新框架增加中国共产党成立和中国特色社会主义新时代等内容，并把香港史置于中国与世界的关联中，"
    "有助于补足近现代国家发展与香港角色的知识链条。彭文平与夏泉（2024）将语文、历史和公民科教材"
    "概括为文化认同、历史认同和政治认同“三位一体”的体系；李臣之等（2022）与王飞（2022）则强调"
    "课程结构、教材审核、探究学习和教师价值引导。上述研究说明制度与教材建设的目标趋于一致，但目标"
    "一致并不自动保证内容均衡和课堂吸收。",
)
set_text(
    paras[11],
    "教材研究呈现出值得并置的不同解释。Vickers（2024）比较2020年通识科与2022年公民科《新视野》"
    "教材，认为历史身份、一国两制、文化与公民权利四个维度均出现明显重构；Xu（2024）比较内地与香港"
    "初中中国历史教材，指出两地长期存在“一国两叙”，香港教材较多以族群文化作为构成性要素，而内地"
    "教材较多以领土和多民族包容建构国家。严瑶（2026）对《新视野公民与社会发展科》三册教材编码发现，"
    "国家安全内容主要通过历史、经济和文化议题关联嵌入，政治安全和经济安全较突出，科技、网络、生态、"
    "生物等新兴安全领域深度相对有限。三项研究共同提示：教材评价既要审查立场与事实，也要关注叙事选择、"
    "领域均衡、多模态表达和学生的证据推理机会。由于研究均以特定版本文本为主，不宜外推至所有教材或"
    "课堂。",
)
set_text(paras[12], "（四）师资配备与专业化挑战：学科知识、课程理解与教学转化")
set_text(
    paras[14],
    "教师专业能力是课程落实的关键。教联会2024年调查显示，在受访教师中，51%表示对教授修订后的初中"
    "中史课程信心不足，76%担忧课时不足，81%支持中史专科专教（见原稿资料24）。这些数据揭示了受访"
    "群体的专业支持需求，但调查样本与抽样方式仍应在引用时一并说明。",
)
set_text(
    paras[15],
    "教育局自2024/25学年推出“史智承传”专业发展计划，为非历史主修教师提供30小时培训（含本地课程与"
    "广州考察）。短期培训可作为入门支持，却难以替代系统的史学训练、课程设计能力和持续同伴教研。"
    "师资问题也不只存在于历史科：公民科涉及法律、经济、社会、国家安全和当代中国，教师需要能够处理"
    "多学科证据、区分事实与价值判断，并把宏观概念转化为适龄的问题与任务。",
)
set_text(paras[16], "（五）实践环节：内地考察的制度覆盖与体验质量")
set_text(
    paras[17],
    "2025/26学年，教育局提供28个公民科内地考察行程及89个学生内地交流计划行程；公开资料显示，自"
    "全面通关以来参与人数已超过20万。规模扩展为学生接触内地社会、历史与科技创造了条件，但参与人数"
    "属于投入指标，尚不能替代对学习过程和学习结果的评估。",
)
set_text(
    paras[18],
    "杜雨璐（2026）结合政策与校本材料，并对6所香港中学的12名教师及管理者开展目的性抽样访谈。研究"
    "发现，标准化路线与校本需要之间存在适配差异，部分带队教师并非公民科教师，现场教学可能让位于纪律"
    "和安全管理；学生的认知更新、情感联结和身份协商，则取决于行前铺垫、现场追问、同伴互动和行后反思。"
    "这为“走马观花”问题提供了过程性解释，但样本并非随机，结论应理解为典型机制而非全港发生率。"
    "2025/26学年允许教学助理随团，有助于分担部分行政工作，但是否改善教学仍需学校层面的过程证据。",
)

set_text(paras[19], "二、课程变革的成因及其对青少年经验的可能影响")
set_text(
    paras[20],
    "现有困境来自多层因素的叠加：国家与特区层面的课程政策重定向，学校组织与评价制度的激励结构，"
    "教师的学科背景和专业判断，以及青年对本地生活、流动机会与国家发展的具体经验。分析时需要区分"
    "制度原因、实施条件与学生结果，避免把任何单一因素解释为充分原因。",
)
set_text(paras[21], "（一）从通识科到公民科：课程变革的制度与政治背景")
set_text(
    paras[22],
    "Yan与Morris（2025）把通识科的创立与废除放在同一历史框架中考察：2009年创科既回应学制改革、"
    "教师配置和探究学习需要，也借助知识经济、终身学习和国际课程等全球话语取得正当性；2019年后，"
    "国家安全成为课程重组的主导论述，原有协商机制的作用下降。Vickers（2024）进一步把这种变化解释"
    "为从多元探究向统一身份与安全优先的范式转换；李臣之等（2022）和王飞（2022）则从课程纠偏、"
    "知识整合与正向价值引导角度评价改革。报告不预设其中某一解释为唯一答案，而把这些分歧视为需要以"
    "课堂观察、教师访谈和学生作品继续检验的竞争性解释。",
)
set_text(paras[23], "（二）青年社会经济焦虑、地方连结与多重认同")
set_text(
    paras[24],
    "国民教育的长期效果与青少年的身份认同、社会参与和未来预期有关。近年调查呈现的并非单向度图景："
    "青年可能同时对社会流动感到焦虑、对香港保持强烈地方连结，并以学习、就业或社区行动表达参与意愿。"
    "这些变量之间的关系需要纵向资料验证，不能把经济焦虑直接等同于政治疏离。",
)
set_text(
    paras[27],
    "对国民教育而言，以上数据提示课程应把国家、香港与个人发展置于可讨论的具体关系中。若教学只停留"
    "在知识告知和守法要求，较难回应青年对教育、职业、住房和社会参与的真实关切；但把发展机会简单作为"
    "认同教育的工具，同样可能削弱教育的可信度。更稳妥的做法是以真实政策问题、可核查数据和多方经验"
    "组织探究，让学生形成有证据的判断，并在共同完成任务的过程中建立参与感。",
)

boundary_h = insert_after(paras[27], "（三）证据边界：从“意图课程”走向实施与结果评估", "Heading 3")
boundary_p = insert_after(
    boundary_h,
    "本报告所用的大部分学术材料分析课程指引或教材文本。Yan与Morris（2025）、Xu（2024）、Vickers"
    "（2024）、彭俊源（2026）和严瑶（2026）均不能直接回答四百余所中学如何实施课程、学生如何理解"
    "教材以及认同是否变化；杜雨璐（2026）虽进入学校实施层面，仍主要依据教师视角。下一阶段调研应"
    "补入分层学校样本、课堂观察、学生访谈与作品分析，并追踪同一批学生的知识、推理、情感和参与变化，"
    "形成“文件—教材—课堂—学生”的证据链。",
    "Normal",
)

set_text(paras[28], "三、面向实施质量的系统性优化建议")
set_text(
    paras[29],
    "未来优化的重点不宜只是增加内容、活动或行政要求，而应提高课程连贯性、教师专业判断、学习任务质量"
    "和评估证据的可用性。建议采用“小规模试点—独立评估—分阶段扩展”的方式，避免未经验证即全面推行。",
)
set_text(paras[30], "（一）优化课程衔接、评估机制与教材质量")
set_text(
    paras[31],
    "建立跨学段课程衔接审查。由课程发展、学科教学和前线教师组成联合小组，依据“主题连续、概念递进、"
    "程序性知识、高阶认知、情感发展”五项指标审查小学人文科、初中公经社与高中公民科。对“达标/未达标”"
    "评价，可先试行不影响原有减负目标的校本“优秀表现”认证或学习档案，观察其对学习投入、公平性和教师"
    "工作量的影响，再决定是否调整公开考试等级及大学收生权重。",
)
set_text(
    paras[32],
    "提升教材编审的专业与透明度。可把学者参与中小学历史、公民和国安教育教材编审纳入大学知识转移或"
    "社会服务认可，同时建立跨学科审读机制，公开事实核查、图片来源、争议议题处理和版本修订原则。教材"
    "不只要保证政治与法律表述准确，也应提供来源多样、难度适切的材料，让学生练习比较证据、识别叙事"
    "选择和解释不同观点。",
)
set_text(
    paras[33],
    "以能力认证逐步减少无支持的非专业兼教。与其设定单一学历门槛，可在五年内建立分层认证：入门培训"
    "解决基本课程与法律知识，进阶微证书覆盖史学方法、当代中国、议题教学和评估设计，校本导师制与共同"
    "备课支持课堂转化。对持续任教者提供带薪进修和课务减免，并以教学档案、观课和学生任务设计作为认证"
    "依据。",
)
set_text(
    paras[35],
    "从“完成行程”转向“完成学习任务”。教育局可把项目式学习作为资助路线的推荐质量标准，要求学校说明"
    "行前问题、现场证据收集、同伴互动和行后成果，但保留校本选择空间。杜雨璐（2026）的访谈显示，"
    "真正影响体验内化的不是景点数量，而是教师能否围绕学生原有认知进行追问，以及学校是否把考察嵌入"
    "前后课程。",
)
set_text(
    paras[39],
    "合理分担行政工作，保障教师教学时间。学校可利用教学助理和承办机构处理订票、名册、物资与部分流程"
    "协调，但教师仍需承担法定照顾与专业责任。更重要的是在出发前明确“行政负责人、教学负责人、安全"
    "负责人”的分工，并为非公民科带队教师提供简明教学包和现场提问脚本。",
)
set_text(
    paras[41],
    "教育局（标准制定者与证据平台建设者）：整合津贴时保留学校基于学生需要设计路线和任务的空间；建立"
    "抽样式实施评估，综合课程文件、课堂/考察观察、学生作品和匿名反馈，不以活动次数、海报数量或纸面"
    "报告作为主要成效指标。可建设脱敏案例库，公开高质量任务、常见失败模式与改进证据。",
)
set_text(
    paras[43],
    "前线教师（知识组织者、价值引导者与释疑者）：面对学生对内地或香港社会问题的疑问，应使用可靠资料"
    "区分事实、解释和价值判断，允许学生提出问题并要求其以证据负责。理性家国情怀并不以回避复杂现实为"
    "前提；教师的专业性体现在能够把成就、差异与挑战放在历史和比较框架中讨论。",
)

set_text(paras[44], "四、内地升学与跨境生活经验：趋势、机制与政策启示")
set_text(
    paras[45],
    "赴内地升学或就业的香港青年提供了观察跨境学习、生活融入与身份协商的重要窗口，但现有公开资料多为"
    "报名统计、媒体个案或横截面调查，尚不足以把其选择归因于某项课程改革。以下分析因此聚焦可观察趋势"
    "与可检验机制，不把该群体描述为单一、同质或已经完成“认同内化”的样本。",
)
set_text(paras[46], "（一）赴内地升学趋势与可能的务实取向")
set_text(
    paras[47],
    "公开统计显示，2025/26学年完成“文凭试收生计划”报名程序的港生为4,521人，约占当年文凭试考生"
    "总数8%；有关调查显示，在计划非本地升学的学生中，选择内地的比例为43.7%。2026/27学年参与该"
    "计划的内地高校增至165所。上述数据说明升学渠道和选择规模扩大，但报名不等于入学，比例变化也可能"
    "同时受学额、学费、专业供给、就业预期和家庭网络影响。",
)
set_text(
    paras[49],
    "可将部分学生的选择概括为“务实取向”：他们把内地高校、产业和城市视为生涯规划的可选空间，并在"
    "“香港人”与“中国人”之间形成并存、层次不同的身份表达。此处的“务实”是分析概念，不应预设其必然"
    "转化为稳定的政治态度；需要以入学前后追踪访谈检验学习、就业机会与身份变化之间的关系。",
)
set_text(
    paras[50],
    "跨境生活可能通过住宿、支付、社团、课程与同伴交往更新学生对内地的日常知识，也可能带来语言习惯、"
    "教学方式和社会网络方面的适应压力。政策评估应同时记录融入经验与困难，并比较不同家庭背景、学校"
    "类型、城市和专业学生的差异，避免只选取成功个案。",
)
set_text(
    paras[52],
    "这类经验的核心启示是：认同与参与往往在真实关系、共同任务和可感知的发展机会中形成，但教育不能"
    "把利益交换等同于价值认同。政策应同时保障学生获得准确信息、平等选择、充分支持和反思空间。",
)
set_text(
    paras[54],
    "推动“走出去”与“引进来”的双向协同。姊妹学校合作可从参访扩展到联合课题、线上共同课堂、教师"
    "教研和学生长期小组项目，并以成果档案记录学习，而非追求形式上的“学分互认”。两地青年在共同解决"
    "真实问题时形成的平视互动，比一次性展示活动更有利于修正刻板印象。",
)
set_text(
    paras[55],
    "把科技与产业发展转化为真实问题情境。人工智能、天文工程、交通基础设施和新能源汽车等案例可以"
    "增强课程的现实感，但不宜停留在“震撼参观”。学生应同时考察技术原理、产业链、职业路径、伦理风险、"
    "区域差异和香港可能发挥的作用，以证据分析代替单向赞叹，使国家发展议题真正转化为可理解、可讨论、"
    "可参与的学习经验。",
)
set_text(
    paras[56],
    "综上，2024至2026年香港中小学国民教育的主要进展，是跨学段课程、教材审核、教师培训和内地考察"
    "的制度覆盖明显增强；主要挑战，则是政策目标与课堂学习之间仍缺少稳定证据链。下一阶段应把工作重心"
    "从“有没有”转向“是否连贯、是否可教、是否投入、学到了什么”，通过跨学段课程审查、分层师资认证、"
    "任务驱动型考察和多源评估，逐步形成兼顾国家认同、地方经验、批判思考与个人发展的教育体系。",
)

# Reference section: retain the original numbered web/policy list and add the academic sources actually used.
set_text(paras[57], "参考资料")
paras[57].style = "Heading 2"
insert_before(paras[58], "政策、统计、调研与媒体资料（沿用原稿编号）", "Heading 3")
academic_h = insert_after(paras[114], "补充学术文献与学位论文", "Heading 3")
academic_refs = [
    "Vickers, E. (2024). The motherland’s suffocating embrace: schooling and public discourse on Hong Kong identity under the National Security Law. Comparative Education, 60(1), 138–158. https://doi.org/10.1080/03050068.2023.2212351",
    "Xu, S. (2024). ‘One country, two narratives’ of China as a nation: comparing Chinese history education in mainland China and Hong Kong. Comparative Education, 60(4), 573–590.",
    "Yan, K. C. A., & Morris, P. (2025). The life and death of Liberal Studies: explaining curriculum change in post-handover Hong Kong. Journal of Curriculum Studies, 57(4), 480–494.",
    "彭俊源（2026）。《香港中学阶段公民教育课程指引衔接研究——以〈公民与社会发展科课程及评估指引〉〈公民、经济与社会课程指引〉为例》（硕士学位论文）。华南师范大学。",
    "杜雨璐（2026）。《构想、实践与体验：香港公民科内地考察的空间生产研究》（硕士学位论文）。华南师范大学。",
    "严瑶（2026）。《多模态视域下香港“公民与社会发展科”教科书国家安全教育的内容建构研究》（硕士学位论文）。华南师范大学。",
    "李臣之、方丽、梁舒婷（2022）。香港《公民与社会发展科课程及评估指引》的特点与启示。《课程·教材·教法》，42（8）。https://doi.org/10.19877/j.cnki.kcjcjf.2022.08.019",
    "彭文平、夏泉（2024）。香港基础教育革新中教材的中华民族共同体叙事。《暨南学报（哲学社会科学版）》，2024（11），1–15。https://doi.org/10.11778/j.jnxb.20231789",
    "王飞（2022）。香港通识教育科的最新发展与未来展望——基于香港青年学生国家认同的视角。《统一战线学研究》，2022（5），91–100。https://doi.org/10.13946/j.cnki.jcqis.2022.05.008",
]
cursor = academic_h
for ref in academic_refs:
    cursor = insert_after(cursor, ref, "Normal")
    cursor.paragraph_format.left_indent = Cm(0.74)
    cursor.paragraph_format.first_line_indent = Cm(-0.74)
    cursor.paragraph_format.space_after = Pt(5)

# Page and typography system.
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.15)
    footer_p = section.footer.paragraphs[0]
    footer_p.clear()
    add_page_number(footer_p)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
normal.paragraph_format.line_spacing = 1.45
normal.paragraph_format.first_line_indent = Cm(0.74)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.widow_control = True

heading_tokens = {
    "Heading 1": (18, (31, 56, 84), 16, 10),
    "Heading 2": (14, (31, 78, 121), 14, 6),
    "Heading 3": (11.5, (42, 91, 132), 10, 4),
    "Heading 4": (10.5, (42, 91, 132), 8, 3),
}
for name, (size, color, before, after) in heading_tokens.items():
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(*color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.widow_control = True

title = doc.paragraphs[0]
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after = Pt(12)
title.paragraph_format.keep_with_next = True

# References begin on a new page.
for p in doc.paragraphs:
    if p.text == "参考资料":
        p.paragraph_format.page_break_before = True
        break

# Existing web references: compact, hanging layout without disturbing their numbering/hyperlinks.
reference_mode = False
for p in doc.paragraphs:
    if p.text == "政策、统计、调研与媒体资料（沿用原稿编号）":
        reference_mode = True
        continue
    if p.text == "补充学术文献与学位论文":
        reference_mode = False
    if reference_mode and p.style.name == "Normal":
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

# Table geometry and styling.
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.9), Cm(3.1), Cm(5.2), Cm(5.8)]
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            cell.width = widths[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), "D9E8F5" if ri == 0 else ("F4F8FB" if ri % 2 == 0 else "FFFFFF"))
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=8.5, east_asia="Songti SC", bold=(ri == 0))
    set_repeat_table_header(table.rows[0])

# Remove the Google Sans direct formatting that caused missing Chinese glyphs in rendering.
set_all_xml_run_fonts(doc)
for p in doc.paragraphs:
    for run in p.runs:
        if p.style.name.startswith("Heading"):
            set_run_font(run, east_asia="PingFang SC")
        else:
            set_run_font(run, east_asia="Songti SC")

doc.core_properties.title = "香港中小学国民教育发展现状、实施张力与优化建议"
doc.core_properties.subject = "基于政策文本、教材研究与前线教师材料的综合分析"
doc.core_properties.comments = "修订版：补充三篇期刊论文、raw目录相关论文及三篇硕士学位论文。"
doc.save(OUTPUT)
print(OUTPUT)
